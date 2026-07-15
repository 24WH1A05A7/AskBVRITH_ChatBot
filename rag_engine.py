"""
rag_engine.py
�??�??�??�??�??�??�??�??�??�??�??�??�??
RAGEngine — the class imported by app.py.

Wraps LangChain + ChromaDB to provide:
  - Document loading from data/bvrit_knowledge_base.docx
  - HuggingFace embeddings (local, no API key needed)
  - ChromaDB vector store persisted to vectorstore/chroma/
  - OpenRouter LLM via ChatOpenAI-compatible interface
  - Tool-augmented query answering (fee calculator, date checker)
  - Section filtering, source formatting, query caching

API surface expected by app.py
�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
  engine = RAGEngine(chunk_size, chunk_overlap, top_k)
  status = engine.initialize()            → dict
  result = engine.query_with_tools(       → dict
               question, chat_history,
               section_filter, verbose)
  sections = engine.get_sections()        → list[str]
  sources   = engine.format_sources(docs) → str
  img_paths = engine.extract_images_from_docs(docs) → list[str]
  engine.set_prompt_variant(variant)
  engine.llm.model                        → str
  engine.embeddings.model                 → str
"""

from __future__ import annotations

import hashlib
import json
import math
import os
import re
import time
from pathlib import Path
from typing import Any

from dotenv import load_dotenv

load_dotenv()

# �??�?? Paths �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
ROOT = Path(__file__).parent
DATA_DIR = ROOT / "data"
VECTORSTORE_DIR = ROOT / "vectorstore" / "chroma"
KNOWLEDGE_BASE = DATA_DIR / "bvrit_knowledge_base.docx"

# Ensure directories exist
VECTORSTORE_DIR.mkdir(parents=True, exist_ok=True)

# �??�?? Config from environment �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE_URL = os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "openai/gpt-oss-20b:free")
EMBEDDING_MODEL = os.getenv(
    "HUGGINGFACE_EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2"
)
TEMPERATURE = float(os.getenv("TEMPERATURE", "0.2"))
COLLEGE_PHONE = os.getenv("COLLEGE_PHONE", "+91-40-27264101")
COLLEGE_EMAIL = os.getenv("COLLEGE_EMAIL", "info@bvrit.ac.in")
COLLEGE_WEBSITE = os.getenv("COLLEGE_WEBSITE", "https://www.bvrit.ac.in")

# Fallback model chain (all free on OpenRouter, verified working)
_DEFAULT_FALLBACKS = [
    {"provider": "openrouter", "model": "openai/gpt-oss-20b:free"},
    {"provider": "openrouter", "model": "nvidia/nemotron-nano-9b-v2:free"},
    {"provider": "openrouter", "model": "nvidia/nemotron-3-nano-30b-a3b:free"},
    {"provider": "openrouter", "model": "nvidia/nemotron-3-super-120b-a12b:free"},
    {"provider": "openrouter", "model": "google/gemma-4-26b-a4b-it:free"},
    {"provider": "openrouter", "model": "poolside/laguna-xs-2.1:free"},
    {"provider": "openrouter", "model": "openrouter/free"},
]

_raw_fallback = os.getenv("FALLBACK_MODELS", "")
try:
    FALLBACK_MODELS: list[dict] = json.loads(_raw_fallback) if _raw_fallback.strip().startswith("[") else _DEFAULT_FALLBACKS
except Exception:
    FALLBACK_MODELS = _DEFAULT_FALLBACKS


# �??�?? Tiny LLM wrapper so engine.llm.model works �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
class _LLMHandle:
    """Thin holder so app.py can read engine.llm.model."""

    def __init__(self, model: str) -> None:
        self.model = model


class _EmbeddingsHandle:
    """Thin holder so app.py can read engine.embeddings.model."""

    def __init__(self, model: str) -> None:
        self.model = model


# �??�?? Document loading �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
def _load_docx(path: Path) -> tuple[list, dict]:
    from docx import Document as DocxDocument
    from langchain_core.documents import Document

    doc = DocxDocument(str(path))
    documents: list[Document] = []
    current_section = "General"
    page = 1
    para_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or style_name == "Title":
            current_section = text
            para_idx = 0
            continue
        para_idx += 1
        if para_idx > 0 and para_idx % 25 == 0:
            page += 1
        documents.append(
            Document(
                page_content=text,
                metadata={"filename": path.name, "page": page, "section": current_section, "source": str(path)},
            )
        )

    for t_idx, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            documents.append(
                Document(
                    page_content="\n".join(rows),
                    metadata={"filename": path.name, "page": page, "section": f"Table {t_idx + 1}", "source": str(path)},
                )
            )

    file_hash = hashlib.sha256(path.read_bytes()).hexdigest()
    meta = {"filename": path.name, "file_hash": file_hash, "document_count": len(documents)}
    return documents, meta


def _split_documents(docs: list, chunk_size: int, chunk_overlap: int) -> list:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)


# �??�?? Vector store �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
def _get_embeddings():
    from langchain_community.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)


def _meta_path(chunk_size: int) -> Path:
    return VECTORSTORE_DIR / f"index_meta_{chunk_size}.json"


def _needs_rebuild(file_hash: str, chunk_size: int) -> bool:
    mp = _meta_path(chunk_size)
    if not mp.exists():
        return True
    try:
        meta = json.loads(mp.read_text(encoding="utf-8"))
        return meta.get("file_hash") != file_hash
    except Exception:
        return True


def _save_meta(chunk_size: int, file_hash: str, chunk_count: int) -> None:
    _meta_path(chunk_size).write_text(
        json.dumps({"file_hash": file_hash, "chunk_count": chunk_count, "chunk_size": chunk_size}, indent=2),
        encoding="utf-8",
    )


# �??�?? LLM factory �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
def _build_openrouter_llm(model: str, temperature: float = TEMPERATURE, streaming: bool = False):
    from langchain_openai import ChatOpenAI

    return ChatOpenAI(
        model=model,
        temperature=temperature,
        api_key=OPENROUTER_API_KEY,
        base_url=OPENROUTER_BASE_URL,
        streaming=streaming,
        default_headers={
            "HTTP-Referer": COLLEGE_WEBSITE,
            "X-Title": "AskBVRITH",
        },
    )


def _is_retryable(exc: BaseException) -> bool:
    msg = str(exc).lower()
    return any(k in msg for k in ("429", "503", "rate limit", "rate_limit", "overloaded", "temporarily unavailable"))


# �??�?? Validation helpers �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??

# Known department/branch aliases used for contradiction detection
_DEPT_ALIASES: dict[str, str] = {
    "cse": "CSE", "computer science": "CSE", "computer": "CSE",
    "ece": "ECE", "electronics": "ECE", "electronics and communication": "ECE",
    "eee": "EEE", "electrical": "EEE", "electrical and electronics": "EEE",
    "it": "IT", "information technology": "IT",
    "mech": "MECH", "mechanical": "MECH",
    "civil": "CIVIL",
    "ai&ml": "AI&ML", "aiml": "AI&ML", "artificial intelligence": "AI&ML",
    "data science": "DS", "ds": "DS",
}

# Keywords that indicate the query is clearly unrelated to BVRIT/college topics
_GENERIC_OOS_KEYWORDS = [
    "fifa", "world cup", "cricket", "ipl", "football", "soccer",
    "stock market", "share price", "recipe", "cooking", "movie", "netflix",
    "weather", "celebrity", "political party", "election", "war", "military",
    "cryptocurrency", "bitcoin", "forex", "astrology", "horoscope",
    "weight loss", "diet plan", "workout routine",
]


def _detect_dept_contradictions(text: str) -> list[str]:
    """
    Return the list of canonical department names found in *text*.
    If more than one distinct department is mentioned it signals a contradiction.
    """
    t = text.lower()
    found: dict[str, str] = {}  # canonical_name → matched keyword
    for keyword, canonical in _DEPT_ALIASES.items():
        # Use word-boundary matching to avoid false positives (e.g. "it" inside "it's")
        pattern = r"\b" + re.escape(keyword) + r"\b"
        if re.search(pattern, t):
            found[canonical] = keyword
    return list(found.keys())


def _validate_fee_request(text: str) -> str | None:
    """
    Pre-flight validation for fee / scholarship queries.

    Returns a short, user-facing error string if the request is invalid,
    or None if the request should proceed normally.
    """
    t_lower = text.lower()

    # �??�?? 1. Contradictory department names �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
    depts = _detect_dept_contradictions(t_lower)
    if len(depts) > 1:
        return (
            f"Your request mentions multiple departments ({', '.join(depts)}). "
            "Please specify a single department to calculate fees."
        )

    # �??�?? 2. Scholarship percentage validation �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
    # Pattern A: number (with optional % or "percent") immediately before "scholarship"
    # Pattern B: "scholarship" followed by up to 3 optional linking words then a number
    #            (covers "scholarship is -20", "scholarship of -20%", "scholarship for -20 percent")
    scholarship_match = re.search(
        r"(-?\d+(?:\.\d+)?)\s*(?:%|percent)?\s*scholarship", t_lower
    ) or re.search(
        r"scholarship\s+(?:\w+\s+){0,3}?(-?\d+(?:\.\d+)?)\s*(?:%|percent)?(?!\s*\w)", t_lower
    )
    if scholarship_match:
        pct = float(scholarship_match.group(1))
        if pct < 0 or pct > 100:
            return (
                "Scholarship percentage must be between 0% and 100%. "
                "Please provide a valid percentage."
            )

    # �??�?? 3. Duration / year validation �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
    duration_match = re.search(
        r"(?:for\s+)?(\d+)\s*(?:-\s*)?year(?:s)?", t_lower
    )
    if duration_match:
        years = int(duration_match.group(1))
        if years == 0:
            return "Program duration must be between 1 and 4 years."
        if years > 4:
            return f"B.Tech at BVRIT is a 4-year program. Duration must be between 1 and 4 years."

    return None  # All checks passed


def _is_generic_out_of_scope(text: str) -> bool:
    """Return True when the query is clearly unrelated to college / BVRIT."""
    t = text.lower()
    return any(kw in t for kw in _GENERIC_OOS_KEYWORDS)


# �??�?? Tools �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
def _run_fee_calculator(branch: str = "", year: str = "", category: str = "general") -> str:
    fees = {
        "cse": {"general": 120000, "management": 150000},
        "ece": {"general": 110000, "management": 140000},
        "eee": {"general": 105000, "management": 135000},
        "it":  {"general": 115000, "management": 145000},
        "mech": {"general": 100000, "management": 130000},
        "civil": {"general": 95000, "management": 125000},
    }
    b = branch.lower().strip()
    c = category.lower().strip()
    fee_info = fees.get(b, fees.get("cse", {}))
    amount = fee_info.get(c, fee_info.get("general", 120000))
    return f"Annual tuition fee for {branch.upper() or 'CSE'} ({category}): ₹{amount:,}/year. Total 4-year: ₹{amount * 4:,}"


def _run_percentage_calculator(obtained: float = 0, total: float = 100) -> str:
    if total == 0:
        return "Error: total marks cannot be zero."
    pct = (obtained / total) * 100
    return f"Percentage: {pct:.2f}% ({obtained}/{total})"


def _run_date_checker(event: str = "") -> str:
    import datetime
    today = datetime.date.today()
    events = {
        "admission": datetime.date(today.year, 6, 30),
        "exam": datetime.date(today.year, 11, 15),
        "result": datetime.date(today.year, 12, 20),
        "counselling": datetime.date(today.year, 7, 15),
    }
    e = event.lower()
    for key, date in events.items():
        if key in e:
            delta = (date - today).days
            if delta < 0:
                return f"{event} was {abs(delta)} days ago ({date.strftime('%d %b %Y')})."
            return f"{event} is in {delta} days — on {date.strftime('%d %b %Y')}."
    return f"Today is {today.strftime('%d %b %Y')}. Please check bvrit.ac.in for event dates."


TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "fee_calculator",
            "description": "Calculate BVRIT tuition fees for a branch and category",
            "parameters": {
                "type": "object",
                "properties": {
                    "branch": {"type": "string", "description": "Branch code e.g. CSE, ECE, EEE"},
                    "year": {"type": "string", "description": "Year of study (1-4)"},
                    "category": {"type": "string", "description": "Admission category: general or management"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "percentage_calculator",
            "description": "Calculate percentage from marks obtained and total marks",
            "parameters": {
                "type": "object",
                "properties": {
                    "obtained": {"type": "number", "description": "Marks obtained"},
                    "total": {"type": "number", "description": "Total marks"},
                },
                "required": ["obtained", "total"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "date_checker",
            "description": "Check days remaining for BVRIT events like admission, exam, result",
            "parameters": {
                "type": "object",
                "properties": {
                    "event": {"type": "string", "description": "Event name e.g. admission, exam, result"},
                },
                "required": ["event"],
            },
        },
    },
]

from langchain_core.tools import tool as _lc_tool  # noqa: E402 (after imports settle)


def _validate_tool_inputs(name: str, args: dict) -> str | None:
    """
    Guardrail layer: validate all tool arguments before execution.
    Returns an error string if invalid, None if OK.
    """
    if name == "fee_calculator":
        # Branch must be a known code
        branch = args.get("branch", "").lower().strip()
        valid_branches = {"cse", "ece", "eee", "it", "mech", "civil", "ai&ml", "ds"}
        if branch and branch not in valid_branches:
            return f"'{args.get('branch')}' is not a recognised branch. Valid branches: CSE, ECE, EEE, IT, MECH, CIVIL."

        # Year must be 1�??4
        year_raw = str(args.get("year", "")).strip()
        if year_raw:
            try:
                year_val = int(year_raw)
                if year_val < 1 or year_val > 4:
                    return "Year of study must be between 1 and 4."
            except ValueError:
                return "Year of study must be a number between 1 and 4."

        # Category must be general or management
        category = args.get("category", "general").lower().strip()
        if category not in ("general", "management", ""):
            return f"Category must be 'general' or 'management', not '{args.get('category')}'."

        # Scholarship embedded in args (future-proof)
        scholarship = args.get("scholarship_pct")
        if scholarship is not None:
            try:
                s = float(scholarship)
                if s < 0 or s > 100:
                    return "Scholarship must be between 0% and 100%."
            except (ValueError, TypeError):
                return "Scholarship must be a numeric percentage between 0 and 100."

    elif name == "percentage_calculator":
        obtained = args.get("obtained")
        total = args.get("total")
        try:
            if float(total) == 0:
                return "Total marks cannot be zero."
            if float(obtained) < 0:
                return "Marks obtained cannot be negative."
            if float(obtained) > float(total):
                return "Marks obtained cannot exceed total marks."
        except (ValueError, TypeError):
            return "Marks obtained and total must be valid numbers."

    elif name == "date_checker":
        event = args.get("event", "")
        if not event or not isinstance(event, str) or len(event.strip()) < 2:
            return "Please provide a valid event name (e.g. admission, exam, result, counselling)."

    else:
        return f"Unsupported operation '{name}'. Available tools: fee_calculator, percentage_calculator, date_checker."

    return None  # All checks passed


def _execute_tool(name: str, args: dict) -> str:
    # Run guardrail before executing
    validation_error = _validate_tool_inputs(name, args)
    if validation_error:
        return f"⚠️ Validation error: {validation_error}"

    if name == "fee_calculator":
        return _run_fee_calculator(**{k: v for k, v in args.items() if k in ("branch", "year", "category")})
    if name == "percentage_calculator":
        return _run_percentage_calculator(**args)
    if name == "date_checker":
        return _run_date_checker(**args)
    return f"Unknown tool: {name}"


# �??�?? System prompt �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
_SYSTEM_PROMPT_A = """\
You are AskBVRITH, the official AI assistant for BVRIT Hyderabad College of Engineering for Women.
Answer questions accurately using ONLY the provided context. If you don't find relevant information
in the context, say so honestly rather than guessing.

Rules:
- Be helpful, concise, and professional.
- Cite sources as [Section, Page N] when referencing context.
- For calculations, use the provided tools.
- Do not invent facts not present in the context.
- If the question is about BVRIT/the college but the answer is not in the context, acknowledge you
  don't have that information and direct the student to contact BVRIT: {contact}
- If the question is completely unrelated to BVRIT or college matters, politely state it is outside
  your scope — do NOT direct unrelated queries to BVRIT contact details.

Context:
{context}
"""

_SYSTEM_PROMPT_B = """\
You are AskBVRITH, a friendly and knowledgeable assistant for BVRIT Hyderabad College.
Use the context below to answer accurately. Be warm, clear, and direct.

Important:
- Only use information from the context provided.
- Cite relevant sections: [Section, Page N].
- For fee or date calculations, use the available tools.
- If the question is about BVRIT/the college but the answer is not in the context, acknowledge you
  don't have that information and direct the student to {contact}.
- If the question is completely unrelated to BVRIT or college matters, politely state it is outside
  your scope — do NOT direct unrelated queries to BVRIT contact details.

Context:
{context}
"""


# �??�?? Main RAGEngine class �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
class RAGEngine:
    """
    RAGEngine wraps ChromaDB + OpenRouter LLM.

    Parameters mirror app.py's sliders so the user can tune them in real time.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        top_k: int = 8,
    ) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.top_k = top_k
        self._prompt_variant = "a"

        # Resolved at initialize() time
        self._vectorstore = None
        self._embeddings_obj = None
        self._sections: list[str] = ["All Sections"]

        model_name = OPENROUTER_MODEL
        self.llm = _LLMHandle(model_name)
        self.embeddings = _EmbeddingsHandle(EMBEDDING_MODEL)

    # �??�?? Public API �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??

    def initialize(self) -> dict:
        """Build or load the vector store. Returns a status dict for the sidebar."""
        from langchain_chroma import Chroma

        if not KNOWLEDGE_BASE.exists():
            raise FileNotFoundError(
                f"Knowledge base not found at {KNOWLEDGE_BASE}. "
                "Please place bvrit_knowledge_base.docx in the data/ folder."
            )

        raw_docs, kb_meta = _load_docx(KNOWLEDGE_BASE)
        file_hash = kb_meta["file_hash"]

        emb = _get_embeddings()
        self._embeddings_obj = emb
        coll_name = f"bvrit_kb_{self.chunk_size}"

        if _needs_rebuild(file_hash, self.chunk_size):
            chunks = _split_documents(raw_docs, self.chunk_size, self.chunk_overlap)
            self._vectorstore = Chroma.from_documents(
                documents=chunks,
                embedding=emb,
                collection_name=coll_name,
                persist_directory=str(VECTORSTORE_DIR),
            )
            _save_meta(self.chunk_size, file_hash, len(chunks))
            index_status = "Newly built"
            chunk_count = len(chunks)
        else:
            self._vectorstore = Chroma(
                collection_name=coll_name,
                embedding_function=emb,
                persist_directory=str(VECTORSTORE_DIR),
            )
            meta = json.loads(_meta_path(self.chunk_size).read_text(encoding="utf-8"))
            index_status = "Loaded from disk"
            chunk_count = meta.get("chunk_count", 0)

        # Populate sections
        try:
            data = self._vectorstore.get()
            sections = sorted({m.get("section", "General") for m in data.get("metadatas", [])})
            self._sections = ["All Sections"] + sections
        except Exception:
            self._sections = ["All Sections"]

        return {
            "document": kb_meta["filename"],
            "chunk_count": chunk_count,
            "index_status": index_status,
        }

    def set_prompt_variant(self, variant: str | None) -> None:
        if variant:
            self._prompt_variant = str(variant).lower()

    def get_sections(self) -> list[str]:
        return self._sections or ["All Sections"]

    def format_sources(self, docs: list) -> str:
        if not docs:
            return ""
        parts: list[str] = []
        seen: set[str] = set()
        for doc in docs:
            section = doc.metadata.get("section", "General")
            page = doc.metadata.get("page", "?")
            key = f"{section} — p.{page}"
            if key not in seen:
                seen.add(key)
                parts.append(key)
        return " · ".join(parts)

    def extract_images_from_docs(self, docs: list) -> list[str]:
        # No image extraction from docx in this implementation
        return []

    def query_with_tools(
        self,
        question: str,
        chat_history: list,
        section_filter: str = "All Sections",
        verbose: bool = False,
    ) -> dict[str, Any]:
        """Run RAG + optional tool calling. Returns dict compatible with app.py."""
        from langchain_core.messages import AIMessage, HumanMessage, SystemMessage, ToolMessage

        start = time.perf_counter()

        # �??�?? 0. Pre-flight guardrail checks �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
        contact = f"Phone {COLLEGE_PHONE}, Email {COLLEGE_EMAIL}, Website {COLLEGE_WEBSITE}"

        # 0a. Generic out-of-scope (completely unrelated to college)
        if _is_generic_out_of_scope(question):
            return {
                "answer": (
                    "I'm AskBVRITH, your BVRIT Hyderabad College assistant. "
                    "That question is outside my scope — I can only help with college-related topics "
                    "such as admissions, fees, departments, placements, and campus facilities."
                ),
                "source_documents": [],
                "latency": round(time.perf_counter() - start, 2),
                "routing": "guardrail_oos",
                "tool_calls": [],
            }

        # 0b. Fee/scholarship/duration/contradiction validation
        fee_keywords = ("fee", "fees", "scholarship", "tuition", "cost", "calculate", "year", "duration")
        if any(kw in question.lower() for kw in fee_keywords):
            validation_error = _validate_fee_request(question)
            if validation_error:
                return {
                    "answer": f"⚠️ {validation_error}",
                    "source_documents": [],
                    "latency": round(time.perf_counter() - start, 2),
                    "routing": "guardrail_validation",
                    "tool_calls": [],
                }

        # 1. Rewrite the question into a self-contained search query using recent history.
        #    This resolves follow-ups like "tell me more about the first one" into explicit queries.
        retrieval_query = self._rewrite_query(question, chat_history)

        # 2. Retrieve context using the rewritten query
        docs = self._retrieve(retrieval_query, section_filter)
        context = self._format_context(docs)

        # 2. Build system prompt
        template = _SYSTEM_PROMPT_A if self._prompt_variant != "b" else _SYSTEM_PROMPT_B
        system_content = template.format(context=context, contact=contact)

        # 3. Assemble message list
        messages: list = [SystemMessage(content=system_content)]
        for msg in (chat_history or [])[-20:]:
            if hasattr(msg, "type"):
                messages.append(msg)
            elif isinstance(msg, dict):
                role = msg.get("role", "")
                content = msg.get("content", "")
                if role == "user":
                    messages.append(HumanMessage(content=content))
                elif role == "assistant":
                    messages.append(AIMessage(content=content))
        messages.append(HumanMessage(content=question))

        # 4. Call LLM with fallback chain + tool loop
        answer, tool_log, used_model, routing = self._llm_with_tools(messages)

        # 5. Post-process
        answer = _format_citations(answer)
        latency = round(time.perf_counter() - start, 2)
        self.llm.model = used_model  # update so observability can read it

        return {
            "answer": answer,
            "source_documents": docs,
            "latency": latency,
            "routing": routing,
            "tool_calls": tool_log,
        }

    # �??�?? Internal helpers �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??

    def _rewrite_query(self, question: str, chat_history: list) -> str:
        """
        Rewrite the user's question into a standalone, keyword-rich search query.
        - For follow-up questions: resolves pronouns/references using chat history.
        - For all questions: expands into domain-specific terms that match the KB.
        Falls back to the original question if rewriting fails.
        """
        # Build a short summary of recent exchanges for context
        recent = []
        for msg in (chat_history or [])[-6:]:
            if isinstance(msg, dict):
                role = msg.get("role", "")
                content = msg.get("content", "")[:200]
            elif hasattr(msg, "type"):
                role = getattr(msg, "type", "")
                content = getattr(msg, "content", "")[:200]
            else:
                continue
            if role in ("user", "human"):
                recent.append(f"User: {content}")
            elif role in ("assistant", "ai"):
                recent.append(f"Assistant: {content}")

        history_text = "\n".join(recent) if recent else "None"

        rewrite_prompt = (
            f"You are a search query optimizer for a college information chatbot.\n"
            f"Conversation so far:\n{history_text}\n\n"
            f"User's question: \"{question}\"\n\n"
            f"Rewrite this as a concise, keyword-rich search query that will find the right "
            f"information in a college knowledge base. "
            f"- Resolve any pronouns or vague references (e.g. 'first', 'it', 'that') using the conversation history.\n"
            f"- Expand abbreviations and use full names (e.g. 'departments' -> 'B.Tech undergraduate branches CSE ECE EEE IT').\n"
            f"- Return ONLY the rewritten query, nothing else."
        )

        try:
            llm = _build_openrouter_llm(FALLBACK_MODELS[0]["model"], temperature=0.0)
            from langchain_core.messages import HumanMessage as HM
            response = llm.invoke([HM(content=rewrite_prompt)])
            rewritten = response.content.strip().strip('"').strip("'")
            if rewritten and len(rewritten) > 5:
                return rewritten
        except Exception:
            pass

        return question

    def _retrieve(self, query: str, section_filter: str = "All Sections") -> list:
        if self._vectorstore is None:
            return []
        try:
            where: dict = {}
            if section_filter and section_filter != "All Sections":
                where["section"] = section_filter

            if where:
                try:
                    results = self._vectorstore.similarity_search_with_score(
                        query, k=self.top_k, filter=where
                    )
                except Exception:
                    results = self._vectorstore.similarity_search_with_score(query, k=self.top_k)
                    results = [(d, s) for d, s in results if d.metadata.get("section") == section_filter]
            else:
                results = self._vectorstore.similarity_search_with_score(query, k=self.top_k)

            # If the query mentions a specific department, also pull all chunks for that
            # department's faculty section directly (they score poorly due to being short names).
            dept_map = {
                "cse": "Computer Science & Engineering (CSE) — 49 faculty",
                "computer science": "Computer Science & Engineering (CSE) — 49 faculty",
                "ece": "Electronics & Communication Engineering (ECE) — 27 faculty",
                "electronics": "Electronics & Communication Engineering (ECE) — 27 faculty",
                "eee": "Electrical & Electronics Engineering (EEE) — 16 faculty",
                "electrical": "Electrical & Electronics Engineering (EEE) — 16 faculty",
                "it": "Information Technology (IT) — 9 faculty",
                "information technology": "Information Technology (IT) — 9 faculty",
                "ai&ml": "CSE — Artificial Intelligence & Machine Learning (AI&ML) — 15 faculty",
                "aiml": "CSE — Artificial Intelligence & Machine Learning (AI&ML) — 15 faculty",
                "artificial intelligence": "CSE — Artificial Intelligence & Machine Learning (AI&ML) — 15 faculty",
                "bsh": "Basic Sciences & Humanities (BS&H) — 33 faculty",
                "basic sciences": "Basic Sciences & Humanities (BS&H) — 33 faculty",
            }
            q_lower = query.lower()
            extra_docs: list = []
            seen_sections: set = set()
            for keyword, target_section in dept_map.items():
                if keyword in q_lower and target_section not in seen_sections:
                    seen_sections.add(target_section)
                    try:
                        extra = self._vectorstore.get(where={"section": target_section})
                        from langchain_core.documents import Document
                        for content, meta in zip(extra.get("documents", []), extra.get("metadatas", [])):
                            doc = Document(page_content=content, metadata=meta)
                            doc.metadata["retrieval_score"] = 0.5  # treat as highly relevant
                            extra_docs.append(doc)
                    except Exception:
                        pass

            # Filter by relevance (L2 distance threshold)
            # Raised from 1.6 to 2.0 �?? short chunks (e.g. single faculty names or
            # branch names) have inherently higher L2 distances even when relevant.
            threshold = 2.0
            docs = []
            seen_contents: set = set()
            for doc, score in results:
                doc.metadata["retrieval_score"] = float(score)
                if score <= threshold and doc.page_content not in seen_contents:
                    seen_contents.add(doc.page_content)
                    docs.append(doc)

            # Append targeted faculty docs (deduplicated)
            for doc in extra_docs:
                if doc.page_content not in seen_contents:
                    seen_contents.add(doc.page_content)
                    docs.append(doc)

            return docs
        except Exception:
            return []

    def _format_context(self, docs: list) -> str:
        if not docs:
            return (
                "NO_CONTEXT_FOUND: No relevant information was found in the BVRIT knowledge base "
                "for this query. If the question is about BVRIT/the college, acknowledge that "
                "you don't have this specific information and provide the official contact details. "
                "Do NOT provide contact details for queries unrelated to the college."
            )
        blocks: list[str] = []
        for doc in docs:
            section = doc.metadata.get("section", "General")
            page = doc.metadata.get("page", "?")
            blocks.append(f"[{section} | Page {page}]\n{doc.page_content}")
        return "\n\n---\n\n".join(blocks)

    def _llm_with_tools(
        self, messages: list
    ) -> tuple[str, list[dict], str, str]:
        """Try the fallback model chain, run tool loop if needed.

        Returns (answer, tool_log, model_used, routing_label).
        """
        from langchain_core.messages import AIMessage, ToolMessage

        tool_log: list[dict] = []
        routing = "rag"
        last_error: str = ""

        for candidate in FALLBACK_MODELS:
            model_name = candidate.get("model", OPENROUTER_MODEL)
            try:
                llm = _build_openrouter_llm(model_name)

                # Some free models don't support tool/function calling.
                # Try with tools first; fall back to plain invocation if it fails.
                try:
                    llm_with_tools = llm.bind_tools(TOOL_DEFINITIONS)
                except Exception:
                    llm_with_tools = llm

                # Tool loop (max 5 iterations)
                msgs = list(messages)
                for _ in range(5):
                    response = llm_with_tools.invoke(msgs)
                    if not isinstance(response, AIMessage) or not response.tool_calls:
                        answer = response.content if hasattr(response, "content") else str(response)
                        return answer, tool_log, model_name, routing

                    # Execute tool calls
                    msgs.append(response)
                    routing = "tool"
                    for tc in response.tool_calls:
                        t_start = time.perf_counter()
                        result = _execute_tool(tc["name"], tc.get("args", {}))
                        tool_log.append({
                            "tool": tc["name"],
                            "args": tc.get("args", {}),
                            "result": result,
                            "latency_ms": (time.perf_counter() - t_start) * 1000,
                        })
                        msgs.append(ToolMessage(content=str(result), tool_call_id=tc["id"]))

                # Exhausted iterations
                return "I completed the tool calculations. Please see the results above.", tool_log, model_name, routing

            except Exception as exc:
                last_error = str(exc)
                # Always try the next model in the chain, regardless of error type.
                # This handles invalid API keys, unsupported models, quota exhaustion, etc.
                continue

        # All models exhausted �?? return a friendly message
        return (
            f"⚠️ All AI models are temporarily unavailable. Please try again shortly. "
            f"For urgent queries contact BVRIT at {COLLEGE_PHONE} or {COLLEGE_EMAIL}.",
            tool_log,
            OPENROUTER_MODEL,
            "error",
        )


# �??�?? Citation formatting �??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??�??
def _format_citations(text: str) -> str:
    """Normalize [Section | Page X] �?? **[Section, Page X]**."""
    pattern = r'\[([^\]|,\n]+)(?:\||,)\s*Page\s*([^\]\n]+)\]'

    def replace(m: re.Match) -> str:
        return f"**[{m.group(1).strip()}, Page {m.group(2).strip()}]**"

    cleaned = re.sub(pattern, replace, text)
    cleaned = cleaned.replace("******", "**").replace("****", "**")
    return cleaned


